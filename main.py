# -*- coding: utf-8 -*-
import docx
import telebot
import requests
from telebot import types
from bs4 import BeautifulSoup

from config import token

bot = telebot.TeleBot(token)
@bot.message_handler(commands=['start'])
def start_message(message):
    global base
    baseFile = open('/data/base.txt', 'r+')
    #baseFile = open('base.txt', 'r+')
    base = baseFile.read().splitlines() 
    if not(str(message.chat.id) in base):
        baseFile.write(str(message.chat.id)+"\n")
    baseFile.close()
    menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
    freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
    orderRefButton = types.KeyboardButton("МФЦ")
    administrationButton = types.KeyboardButton("Руководство института")
    brochureButton = types.KeyboardButton("Брошюра об институте")
    studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
    applicantButton = types.KeyboardButton("Абитуриентам 🤵")
    mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
    newsButton = types.KeyboardButton("Новости 📰")
    contactsButton = types.KeyboardButton("Контакты 👥")
    startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
    CAEButton = types.KeyboardButton("Стратегические академические единицы")
    business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
    menuKeyboard.row(administrationButton)
    menuKeyboard.row(brochureButton)
    menuKeyboard.row(newsButton, contactsButton)
    menuKeyboard.row(freqaskqueButton)
    menuKeyboard.row(orderRefButton)
    menuKeyboard.row(studentsButton, applicantButton)
    menuKeyboard.row(mastersdegreeprogramsButton)
    menuKeyboard.row(startupButton)
    menuKeyboard.row(CAEButton)
    menuKeyboard.row(business_forumButton)
    menuKeyboard.row(constitutionButton)
    mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
    bot.register_next_step_handler(mesg, chooseAction)
@bot.message_handler(commands=['send'])
def checkStatus(message):
    removeKeyboard = types.ReplyKeyboardRemove()
    adminListFile = open('adminList.txt', 'r+')
    adminList = adminListFile.read().splitlines()
    if (str(message.chat.id) in adminList):
        bot.send_message(message.chat.id,f"Колличество пользователей: {len(base)}")
        mesg = bot.send_message(message.chat.id, "Введите сообщение для рассылки:",reply_markup=removeKeyboard)
        bot.register_next_step_handler(mesg, sendMessage)
        adminListFile.close()
    else:
        bot.send_message(message.chat.id,"Действие запрещено",reply_markup=removeKeyboard)
        adminListFile.close()
def sendMessage(message):
    global msg
    msg = message.text
    bot.send_message(message.chat.id,"Сообщение которое будет отправлено:")
    bot.send_message(message.chat.id,msg)
    btns = types.ReplyKeyboardMarkup(resize_keyboard=True)
    plusBtn = types.KeyboardButton("+")
    plusminusBtn = types.KeyboardButton("+-")
    minusBtn = types.KeyboardButton("-")
    btns.row(plusBtn,plusminusBtn,minusBtn)
    mesg = bot.send_message(message.chat.id,"Отправьте:\n'+' для подтверждения отправки\n'+-' для изменения сообщения\n'-' для отмены",reply_markup=btns)
    bot.register_next_step_handler(mesg, confirmation)
def confirmation(message):
    removeKeyboard = types.ReplyKeyboardRemove()
    inputmesg = message.text
    if inputmesg=="+":
        bot.send_message(message.chat.id, "Сообщения отправляются...", reply_markup=removeKeyboard)
        counter = 0
        for user in base:
            try:
                bot.send_message(user, msg)
                counter +=1
            except:
                pass
        bot.send_message(message.chat.id, f"Отправлено {counter} пользователям из {len(base)}")
    elif inputmesg=="+-":
        mesg = bot.send_message(message.chat.id, "Введите новое сообщение:", reply_markup=removeKeyboard)
        bot.register_next_step_handler(mesg, sendMessage)
    elif inputmesg=="-":
        bot.send_message(message.chat.id, "Рассылка отменена", reply_markup=removeKeyboard)
@bot.message_handler(content_types=["text"])
def menu(message):
    menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
    freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
    orderRefButton = types.KeyboardButton("МФЦ")
    administrationButton = types.KeyboardButton("Руководство института")
    brochureButton = types.KeyboardButton("Брошюра об институте")
    studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
    applicantButton = types.KeyboardButton("Абитуриентам 🤵")
    mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
    newsButton = types.KeyboardButton("Новости 📰")
    contactsButton = types.KeyboardButton("Контакты 👥")
    startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
    CAEButton = types.KeyboardButton("Стратегические академические единицы")
    business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
    menuKeyboard.row(administrationButton)
    menuKeyboard.row(brochureButton)
    menuKeyboard.row(newsButton, contactsButton)
    menuKeyboard.row(freqaskqueButton)
    menuKeyboard.row(orderRefButton)
    menuKeyboard.row(studentsButton, applicantButton)
    menuKeyboard.row(mastersdegreeprogramsButton)
    menuKeyboard.row(startupButton)
    menuKeyboard.row(CAEButton)
    menuKeyboard.row(business_forumButton)
    menuKeyboard.row(constitutionButton)
    mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
    bot.register_next_step_handler(mesg, chooseAction)

def chooseAction(message):
    if (message.text == "МФЦ"):
        typeOfRefKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        referenceButton = types.KeyboardButton("Оформить справку 📋")
        defaultRefButton = types.KeyboardButton("↗️")
        refCallButton = types.KeyboardButton("↖️")
        backButton = types.KeyboardButton("Назад")
        typeOfRefKeyboard.row(referenceButton)
        typeOfRefKeyboard.row(defaultRefButton, refCallButton)
        typeOfRefKeyboard.row(backButton)
        bot.send_message(message.chat.id,"Оформление справки происходит через [МФЦ](https://mfc.msal.ru/), для её оформления воспользуйтесь кнопкой:\n«Оформить справку 📋»", parse_mode='MarkdownV2')
        mesg = bot.send_message(message.chat.id,"•Справка об обучении: не требует дополнительных документов.\n\n•Справка-вызов: оформляется только при наличии справки с места работы и заявления; оригинал справки-вызова можно забрать только при предоставлении оригинала справки с места работы и оригинала заявления.",reply_markup=typeOfRefKeyboard)
        bot.register_next_step_handler(mesg, chooseType)
        return
    elif (message.text == "Оформить справку 📋"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference2Button = types.InlineKeyboardButton("Оформить...", url = "https://mfc.msal.ru/")
        newsMarkup.row(reference2Button)
        mesg = bot.send_message(message.chat.id, "👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Руководство института"):
        administrationKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        directorButton = types.KeyboardButton("Директор")
        assistantDirectorButton = types.KeyboardButton("Заместитель директора")
        secretaryButton = types.KeyboardButton("Секретарь")
        scientistButton = types.KeyboardButton("Научный руководитель")
        inspectorsButton = types.KeyboardButton("Инспекторы")
        mainmenuButton = types.KeyboardButton("Главное меню 🗂️")
        administrationKeyboard.row(scientistButton)
        administrationKeyboard.row(directorButton)
        administrationKeyboard.row(assistantDirectorButton)
        administrationKeyboard.row(secretaryButton)
        administrationKeyboard.row(inspectorsButton)
        administrationKeyboard.row(mainmenuButton)
        mesg = bot.send_message(message.chat.id,"Руководство института",reply_markup=administrationKeyboard)
        bot.register_next_step_handler(mesg, choosePerson)
        return

    elif (message.text == "Образовательные программы"):
        programMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        bachelorBtn = types.KeyboardButton("Бакалавриат")
        magistracyBtn = types.KeyboardButton("Магистратура")
        returnToMenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
        programMarkup.row(bachelorBtn, magistracyBtn)
        programMarkup.row(returnToMenuBtn)
        mesg = bot.send_message(message.chat.id,"Выберите программу обучения:", reply_markup=programMarkup)
        bot.register_next_step_handler(mesg, chooseProgram)
        return

    elif (message.text == "Студентам 🧑‍🎓"):
        mesg = bot.send_message(message.chat.id,"[• Календарный учебный график на 2024/2025 учебный год Института бизнес-права](https://view.officeapps.live.com/op/view.aspx?src=https%3A%2F%2Fmsal.ru%2Fupload%2Fmedialibrary%2F821%2F61irfkcdcyg7c7wh5aihhbgoleq090r6.xlsx&wdOrigin=BROWSELINK)\n\n"
                                         "[• Актуальные редакции всех локальных нормативных актов, регламентирующих различные аспекты образовательной деятельности (см. вкладку «Документы по образовательной и методической деятельности»](https://msal.ru/structure/upravleniya/uchebno-metodicheskoe-upravlenie/?ysclid=ld78gcnnds177522123)\n\n"
                                         "[• Подробная информация о всех видах стипендий](https://msal.ru/content/studentam/stipendii/?ysclid=ld781glc19102372321)\n\n"
                                         "[• Все о студенческих общежитиях](https://msal.ru/structure/upravleniya/upravlenie-obshchezhitiy/?ysclid=ld78v4aqg7455516138)\n\n"
                                         "[• Исчерпывающая информация о практиках (см. вкладку «Практическая подготовка»)](https://msal.ru/structure/upravleniya/uchebno-metodicheskoe-upravlenie/tsentr-organizatsii-praktiki-obuchayushchikhsya-i-trudoustroystva-vypusknikov/)\n\n"
                                         "[• Сроки проведения практической подготовки в 2024-2025 учебном году Института бизнес-права](https://disk.yandex.ru/i/kYiQWmYbhie0SA)\n\n"      
                                         "[• Основные правила работы со сносками и библиографией](https://msal.ru/structure/upravleniya/uchebno-metodicheskoe-upravlenie/tsentr-organizatsii-praktiki-obuchayushchikhsya-i-trudoustroystva-vypusknikov/)\n\n"
                                         "[• Учебное пособие для студентов «Пишем курсовую работу по праву»](https://msal.ru/upload/medialibrary/ad3/427jost6ch5jco5lrqyy42mcuvh4p7jc.pdf)\n\n"
                                         "[• Шаблоны слайдов презентации в соответствии с фирменным стилем университета](https://disk.yandex.ru/i/AuWqxGmkd0Yb_Q)\n\n"
                                         "[• График выполнения выпускной квалификационной работы](https://view.officeapps.live.com/op/view.aspx?src=https%3A%2F%2Fwww.msal.ru%2Fupload%2Fmedialibrary%2F275%2F7ktmrjiaz11rruh5kctp107lc5if81r3.docx&wdOrigin=BROWSELINK)\n\n"
                                         "[• Титульные листы магистерских диссертаций](https://disk.yandex.ru/d/_EthuEjcWUYVRw)\n\n"
                                         "[• О военном учебном центре](https://msal.ru/structure/tsentry/voennyy-uchebnyy-tsentr/)\n\n"
                                         "[• Инструкция к онлайн-курсу «История России»](https://msal.ru/upload/medialibrary/5f8/a228hs4pzhd0i76k649yqj2se5miwky6.pdf)\n\n"
                                         "[• Руководство первокурсника Института бизнес-права](https://disk.yandex.ru/i/dSTYZ3YeQF2EgQ)\n\n"
                                         "[• Путеводитель по МГЮА](https://files.msal.ru/HTCOMNET/Handlers/AnonymousDownload.ashx?file=47cc1850)\n\n"
                                         "[• Порядок перехода, перевода и восстановления](https://msal.ru/content/studentam/perevod-i-vosstanovlenie/?ysclid=m2rpscul1j42875846)\n\n"
                                         "[• Информация о кредите на образование с государственной поддержкой (T-банк)](https://www.tbank.ru/loans/cash-loan/education/?utm_source=university&utm_medium=flyer_qr&utm_campaign=offline_form)\n\n"
                                         "[• Информация о кредите на образование с государственной поддержкой (Сбер)](https://www.sberbank.ru/ru/person/credits/money/credit_na_obrazovanie)\n\n"
                                         "[• Положение о курсовом проектировании](https://msal.ru/upload/struktura/upravl/UMU/obr_pr/38.04.04%20-%20Проектная%20деятельность%20в%20государственном%20управлении/2022/Приказ%20№117%20от%2015.03.2024%20Положение%20о%20курсовом%20проектировании%20по%20основным%20профессион_organized.pdf)\n\n", parse_mode='Markdown', disable_web_page_preview=True)
        bot.register_next_step_handler(mesg, chooseAction)
        return

    elif (message.text == "Часто задаваемые вопросы 📄"):
            typeOfRefKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
            ViewthelistButton = types.KeyboardButton("Ознакомиться со списком")
            NotmyqueButton = types.KeyboardButton("Нет ответа на мой вопрос")
            backButton = types.KeyboardButton("Назад 🗂️")
            typeOfRefKeyboard.row(ViewthelistButton)
            typeOfRefKeyboard.row(NotmyqueButton)
            typeOfRefKeyboard.row(backButton)
            mesg = bot.send_message(message.chat.id,"Прежде чем задать вопрос, убедитесь, что он не находится в списке часто задаваемых вопросов 📃",reply_markup=typeOfRefKeyboard)
            bot.register_next_step_handler(mesg, chooseQuestion)
            return
    elif (message.text == "Абитуриентам 🤵"):
        # videosMarkup = types.InlineKeyboardMarkup()
        # legalsupportVideo = types.InlineKeyboardButton("Правовое сопровождение бизнеса", url="https://youtu.be/ANhihH_rOM0")
        # corporatelawVideo = types.InlineKeyboardButton("Корпоративное право",url="https://youtu.be/YMUd_ZD0JL8")
        # videosMarkup.row(legalsupportVideo)
        # videosMarkup.row(corporatelawVideo)
        # mesg =  bot.send_message(message.chat.id, "Видео презентаций магистерских программ 'Правовое сопровождение бизнеса (бизнес-юрист)' и 'Корпоративное право':", reply_markup=videosMarkup)
        # bot.register_next_step_handler(mesg, chooseAction)
        abiturMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        presentationBtn = types.KeyboardButton("Презентация института")
        presentationuniversitBtn = types.KeyboardButton("Презентация университета")
        presentationadmcomBtn = types.KeyboardButton("Презентация приемной комиссии")
        presentationInnjurBtn = types.KeyboardButton("Презентация по инновационной юриспруденции")
        presentationPresentationInsLegTransBtn = types.KeyboardButton("Презентация Института юридического перевода")
        programsBtn = types.KeyboardButton("Магистерские программы")
        faqBtn = types.KeyboardButton("Часто задаваемые вопросы")
        rolikBtn = types.KeyboardButton("Ролик о кампусе на Шитова")
        returnToMenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
        abiturMarkup.row(presentationBtn)
        abiturMarkup.row(presentationuniversitBtn)
        abiturMarkup.row(presentationadmcomBtn)
        abiturMarkup.row(presentationInnjurBtn)
        abiturMarkup.row(presentationPresentationInsLegTransBtn)
        abiturMarkup.row(programsBtn)
        abiturMarkup.row(faqBtn)
        abiturMarkup.row(rolikBtn)
        abiturMarkup.row(returnToMenuBtn)
        mesg = bot.send_message(message.chat.id,"Выберите действие:",reply_markup=abiturMarkup)
        bot.register_next_step_handler(mesg, abiturients)
        return
    elif (message.text == "Новости 📰"):
        newsMarkup = types.InlineKeyboardMarkup()
        vkBtn = types.InlineKeyboardButton("Вконтакте", url = "https://vk.com/businessmsal")
        tgBtn = types.InlineKeyboardButton("Telegram", url = "https://t.me/businessmsal")
        newsMarkup.row(vkBtn, tgBtn)
        mesg = bot.send_message(message.chat.id, "Официальные источники:", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Брошюра об институте"):
        newsMarkup = types.InlineKeyboardMarkup()
        brochBtn = types.InlineKeyboardButton("Брошюра", url = "https://msal.ru/upload/medialibrary/f4d/s6c65nmvzvwml79y4tff3keu4rqx9oyh.pdf")
        newsMarkup.row(brochBtn)
        mesg = bot.send_message(message.chat.id, "👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Контакты 👥"):
        mesg = bot.send_message(message.chat.id,"📬 Почта института:\nbusiness@msal.ru\n\n📞 Телефон института:\n+7 (499) 244-88-88 доб. 798\n\n👨‍💻 Разработчики:\n@ernest_melikyan @Marusik_01 ")
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Стратегические академические единицы"):
        SAEMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        center1 = types.KeyboardButton("Центр правовых исследований банкротства")
        school2 = types.KeyboardButton("Школа предпринимателя")
        backBtn = types.KeyboardButton("Назад")
        SAEMarkup.row(center1, school2)
        SAEMarkup.row(backBtn)
        mesg = bot.send_message(message.chat.id,"Стратегические академические единицы:",reply_markup = SAEMarkup)
        bot.register_next_step_handler(mesg, SAEChoose)
        return
    elif (message.text == "Кутафинский бизнес-форум"):
        forumMarkup = types.InlineKeyboardMarkup()
        forum2022Btn = types.InlineKeyboardButton("I КБФ 2022", url="https://msal.ru/news/-kutafinskiy-biznes-forum-v-mgyua-vspominaem-kak-eto-bylo/")
        forum2023Btn = types.InlineKeyboardButton("II КБФ 2023", url="https://www.msal.ru/news/ii-kutafinskiy-biznes-forum-sostoyalsya-v-mgyua/")
        forum2024Btn = types.InlineKeyboardButton("III КБФ 2024", url="https://msal.ru/news/v-zagorodnom-komplekse-malyushina-dacha-universiteta-imeni-o-e-kutafina-mgyua-sostoyalsya-iii-kutafi/")
        forumMarkup.row(forum2022Btn, forum2023Btn, forum2024Btn)
        mesg = bot.send_message(message.chat.id, "Кутафинский бизнес-форум", reply_markup=forumMarkup)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Программа «Стартап как ВКР»"):
        mesg = bot.send_message(message.chat.id, "Начиная с 2021/2022 учебного года обучающиеся магистратуры вместо традиционного написания выпускной квалификационной работы могут оформить свою предпринимательскую инициативу в форму магистерской диссертации.\n\n"
                                                 "В данный момент программа «Стартап как ВКР» реализуется в качестве эксперимента, поэтому принять участие в ней могут только магистранты магистерских программ «Правовое сопровождение бизнеса (бизнес-юрист)» и «Корпоративное право» которые реализуются кафедрой предпринимательского и корпоративного права на базе Института бизнес-права.\n\n"
                                                 "– Более подробно содержание университетской программы «Стартап как ВКР» [раскрыто в этой презентации.](https://msal.ru/upload/medialibrary/65c/2zu7kr8h9xcx4yj85bv9sns88jg3gkh7.pdf)\n\n"
                                                 "– С полномочиями и составом Экспертного совета по сопровождению программы «Стартап как ВКР», а также информацией об успешно защищенных стартап-проектах можно ознакомиться [по этой ссылке.](https://msal.ru/content/ob-universitete/sovety/ekspertnyy-sovet-po-soprovozhdeniyu-programmy-startap-kak-vkr/)\n\n"
                                                 "– С общей информацией о проекте «Стартап как ВКР» [можно ознакомиться в презентации](https://msal.ru/upload/medialibrary/6bb/l7qogoj3hrx4lpc0l75kkmji85p3pz29.pdf) Министерства науки и высшего образования Российской Федерации.\n\n"
                                                 "– О федеральном проекте «Платформа университетского технологического предпринимательства» [можно узнать здесь.](https://univertechpred.ru/)", parse_mode='Markdown', disable_web_page_preview=True)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    elif (message.text == "Конституция России 🇷🇺"):
        keyboard_constitution = types.ReplyKeyboardMarkup(resize_keyboard=True)
        mainmenu = types.KeyboardButton("Вернуться в меню 🗂️")
        keyboard_constitution.add(mainmenu)
        bot.send_message(message.chat.id, "Введите номер статьи:", reply_markup=keyboard_constitution)
        bot.register_next_step_handler(message, findConstitution)
    elif (message.text == "/send"):
        btns = types.ReplyKeyboardMarkup(resize_keyboard=True)
        continueBtn = types.KeyboardButton("Продолжить")
        btns.row(continueBtn)
        mesg = bot.send_message(message.chat.id,"Нажмите продолжить", reply_markup=btns)
        bot.register_next_step_handler(mesg, checkStatus)
        return
def findConstitution(message):
    if (message.text == "Вернуться в меню 🗂️"):
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)
    else:

        urlconstitution = "https://www.consultant.ru/document/cons_doc_LAW_28399/"
        pageconstitution = requests.get(urlconstitution)
        soup = BeautifulSoup(pageconstitution.text, "html.parser")
        sectionsconstitution = {}
        for i in soup.css.select(".document-page__toc ul ul li a"):
            link = 'https://www.consultant.ru' + i['href']
            sectionsconstitution[i.text] = link
        input = message.text
        for i in sectionsconstitution:
            if (f"Статья {input}" in i):

                # bot.send_message(message.chat.id, i + "\n" + sections.get(i))

                # i --- заголовок статьи: Номер. Название
                # sections.get(i) --- ссылка на номер статьи i
                link = sectionsconstitution.get(i)  # url
                page1 = requests.get(link)
                soupForSite = BeautifulSoup(page1.text, "html.parser")
                for i in (soupForSite.css.select(".document-page__content")):
                    if len(i.text) > 4096:
                        for x in range(0, len(i.text), 4096):
                            bot.send_message(message.chat.id, i.text[x:x + 4096])
                    else:
                        bot.send_message(message.chat.id, i.text)
                break
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        mainmenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
        markup.row(mainmenuBtn)
        bot.send_message(message.chat.id, "Вы можете продолжить вводить номер статьи или выбрать пункт ниже 👇",reply_markup=markup)
        bot.register_next_step_handler(message, findConstitution)

def chooseQuestion(message):
    if (message.text == "Ознакомиться со списком"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference3Button = types.InlineKeyboardButton("Ознакомиться...", url="https://www.msal.ru/structure/instituty/institut-biznes-prava/FAQ/")
        newsMarkup.row(reference3Button)
        bot.send_message(message.chat.id,"👇",reply_markup=newsMarkup)
        bot.register_next_step_handler(message, chooseQuestion)
    elif (message.text == "Нет ответа на мой вопрос"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference4Button = types.InlineKeyboardButton("Задать вопросы...", url = "https://t.me/IBP_SUPPORT")
        newsMarkup.row(reference4Button)
        bot.send_message(message.chat.id,"Если на ваш вопрос отсутствует ответ в списке часто задаваемых вопросах, то задайте его здесь 👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(message, chooseQuestion)
        return
    elif (message.text == "Назад 🗂️"):
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)

def abiturients(message):
    if (message.text == "Презентация института"):
        with open("Презентация Института бизнес-права.pdf", "rb") as file:
            try:
                bot.send_message(message.chat.id,"Файл отправляется...")
                mesg = bot.send_document(message.chat.id, file)
                bot.register_next_step_handler(mesg, abiturients)
            except:
                mesg = bot.send_message(message.chat.id,"Пожалуйста, повторите")
                bot.register_next_step_handler(mesg, abiturients)
                return
    elif (message.text == "Презентация университета"):
        with open("Презентация Университета.pdf", "rb") as file:
            try:
                bot.send_message(message.chat.id,"Файл отправляется...")
                mesg = bot.send_document(message.chat.id, file)
                bot.register_next_step_handler(mesg, abiturients)
            except:
                mesg = bot.send_message(message.chat.id,"Пожалуйста, повторите")
                bot.register_next_step_handler(mesg, abiturients)
                return
    elif (message.text == "Презентация приемной комиссии"):
        with open("Презентация Приемной комиссии.pdf", "rb") as file:
            try:
                bot.send_message(message.chat.id,"Файл отправляется...")
                mesg = bot.send_document(message.chat.id, file)
                bot.register_next_step_handler(mesg, abiturients)
            except:
                mesg = bot.send_message(message.chat.id,"Пожалуйста, повторите")
                bot.register_next_step_handler(mesg, abiturients)
                return
    elif (message.text == "Презентация по инновационной юриспруденции"):
        with open("Презентация Иновационная юриспроденция.pdf", "rb") as file:
            try:
                bot.send_message(message.chat.id,"Файл отправляется...")
                mesg = bot.send_document(message.chat.id, file)
                bot.register_next_step_handler(mesg, abiturients)
            except:
                mesg = bot.send_message(message.chat.id,"Пожалуйста, повторите")
                bot.register_next_step_handler(mesg, abiturients)
                return
    elif (message.text == "Презентация Института юридического перевода"):
        with open("Презентация Института юридического перевода.pdf", "rb") as file:
            try:
                bot.send_message(message.chat.id,"Файл отправляется...")
                mesg = bot.send_document(message.chat.id, file)
                bot.register_next_step_handler(mesg, abiturients)
            except:
                mesg = bot.send_message(message.chat.id,"Пожалуйста, повторите")
                bot.register_next_step_handler(mesg, abiturients)
                return
    elif (message.text == "Магистерские программы"):
        programs2Markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        program1Btn = types.KeyboardButton("Правовое сопровождение бизнеса")
        program2Btn = types.KeyboardButton("Корпоративное право")
        backBtn= types.KeyboardButton("Назад")
        programs2Markup.row(program1Btn,program2Btn)
        programs2Markup.row(backBtn)
        mesg = bot.send_message(message.chat.id, "Магистерские программы:",reply_markup=programs2Markup)
        bot.register_next_step_handler(mesg,Programs2)

    elif (message.text == "Часто задаваемые вопросы"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference6Button = types.InlineKeyboardButton("Ознакомиться...", url="https://www.msal.ru/structure/instituty/institut-biznes-prava/FAQ/")
        newsMarkup.row(reference6Button)
        bot.send_message(message.chat.id, "👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, abiturients)
        return
    elif (message.text == "Ролик о кампусе на Шитова"):
        mesg = bot.send_message(message.chat.id,"Ролик о кампусе на Шитова:\nhttps://vk.com/video-65417_456239327")
        bot.register_next_step_handler(mesg, abiturients)

    elif (message.text == "Вернуться в меню 🗂️"):
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)

def Programs2(message):
    if (message.text == "Правовое сопровождение бизнеса"):
        program1Markup = types.InlineKeyboardMarkup()
        videoBtn = types.InlineKeyboardButton("Видео", url="https://youtu.be/ANhihH_rOM0")
        presentationBtn = types.InlineKeyboardButton("Презентация", url="https://msal.ru/upload/medialibrary/1a4/wy04pbw7l8l8seewzfpl26qjidgwa4gv.pdf")
        siteBtn = types.InlineKeyboardButton("Сайт", url="https://www.msal.ru/programs/vysshee-obrazovanie/magistratura/magisterskaya-programma-pravovoe-soprovozhdenie-biznesa-biznes-yurist/")
        program1Markup.row(videoBtn)
        program1Markup.row(presentationBtn)
        program1Markup.row(siteBtn)
        msg = bot.send_message(message.chat.id, "Правовое сопровождение бизнеса (бизнес-юрист)", reply_markup=program1Markup)
        bot.register_next_step_handler(msg, Programs2)
    elif (message.text == "Корпоративное право"):
        program2Markup = types.InlineKeyboardMarkup()
        videoBtn = types.InlineKeyboardButton("Видео", url="https://youtu.be/YMUd_ZD0JL8?si=PxfJCK0pCsUcH0i_")
        presentationBtn = types.InlineKeyboardButton("Презентация",url="https://msal.ru/upload/medialibrary/848/7lur8yho4ck0w6vdge7jya5n1uk5vb11.pdf")
        siteBtn = types.InlineKeyboardButton("Сайт",url="https://www.msal.ru/programs/vysshee-obrazovanie/magistratura/magisterskaya-programma-korporativnoe-pravo/")
        program2Markup.row(videoBtn)
        program2Markup.row(presentationBtn)
        program2Markup.row(siteBtn)
        msg = bot.send_message(message.chat.id, "Корпоративное право",reply_markup=program2Markup)
        bot.register_next_step_handler(msg, Programs2)
    elif (message.text == "Назад"):
        abiturMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        presentationBtn = types.KeyboardButton("Презентация института")
        presentationuniversitBtn = types.KeyboardButton("Презентация университета")
        presentationadmcomBtn = types.KeyboardButton("Презентация приемной комиссии")
        presentationInnjurBtn = types.KeyboardButton("Презентация по инновационной юриспруденции")
        presentationPresentationInsLegTransBtn = types.KeyboardButton("Презентация Института юридического перевода")
        programsBtn = types.KeyboardButton("Магистерские программы")
        faqBtn = types.KeyboardButton("Часто задаваемые вопросы")
        returnToMenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
        abiturMarkup.row(presentationBtn)
        abiturMarkup.row(presentationuniversitBtn)
        abiturMarkup.row(presentationadmcomBtn)
        abiturMarkup.row(presentationInnjurBtn)
        abiturMarkup.row(presentationPresentationInsLegTransBtn)
        abiturMarkup.row(programsBtn)
        abiturMarkup.row(faqBtn)
        abiturMarkup.row(returnToMenuBtn)
        mesg = bot.send_message(message.chat.id, "Выберите действие:", reply_markup=abiturMarkup)
        bot.register_next_step_handler(mesg, abiturients)
        return


def chooseProgram(message):
    if message.text == "Бакалавриат":
        programmsKeyboard = types.InlineKeyboardMarkup()
        jurisprudenceBtn = types.InlineKeyboardButton("Юриспруденция (на базе ИБП)", url="https://msal.ru/programs/vysshee-obrazovanie/bakalavriat/40-03-01-yurisprudentsiya-na-baze-ibp/")
        programmsKeyboard.row(jurisprudenceBtn)
        mesg = bot.send_message(message.chat.id, "Программа бакалавриата, реализуемая в Институте бизнес-права:",reply_markup=programmsKeyboard)
        bot.register_next_step_handler(mesg, chooseProgram)
    elif message.text == "Магистратура":
        programmsKeyboard = types.InlineKeyboardMarkup()
        legalsupportBtn = types.InlineKeyboardButton("Правовое сопровождение бизнеса", url="https://msal.ru/programs/vysshee-obrazovanie/magistratura/magisterskaya-programma-pravovoe-soprovozhdenie-biznesa-biznes-yurist/")
        corporatelawBtn = types.InlineKeyboardButton("Корпоративное право", url="https://msal.ru/programs/vysshee-obrazovanie/magistratura/magisterskaya-programma-korporativnoe-pravo/")
        programmsKeyboard.row(legalsupportBtn)
        programmsKeyboard.row(corporatelawBtn)
        mesg = bot.send_message(message.chat.id, "Программы магистратуры, реализуемые в Институте бизнес-права:", reply_markup=programmsKeyboard)
        bot.register_next_step_handler(mesg, chooseProgram)
    elif message.text == "Вернуться в меню 🗂️":
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)

def SAEChoose(message):
    returnMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    returnToMenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
    returnMarkup.row(returnToMenuBtn)
    if (message.text == "Центр правовых исследований банкротства"):
        bot.send_message(message.chat.id, "Центр правовых исследований банкротства создан 25 февраля 2021 г. на базе кафедры предпринимательского и корпоративного права, а также Института бизнес-права Университета имени О.Е. Кутафина (МГЮА).\n\n"
                                          "Центр правовых исследований банкротства может быть представлен как экспертная площадка в юридической и образовательной среде, имеющая своей целью междисциплинарное исследование проблем правового регулирования отношений несостоятельности. В работе Центра в настоящее время принимает участие широкий круг заинтересованных «банкротной проблематикой» лиц, включая преподавателей кафедры предпринимательского и корпоративного права, защитивших диссертации по проблемам банкротства (в частности, С.С. Галкин, Е.Е. Енькова, Т.П. Шишмарева), иных работников и преподавателей Университета, специализирующихся по вопросам банкротства, практических работников и обучающихся.\n\n"
                                          "Принять участие в работе Центра будет интересно прежде всего практикующим в сфере банкротства юристам, арбитражным управляющим, обучающимся бакалавриата, магистратуры и аспирантуры, чья область практического и научного интереса связана с проблемами правового регулирования отношений несостоятельности, применения норм законодательства о банкротстве.\n\n"
                                          "Общая информация о планируемой деятельности Центра правовых исследований банкротства изложена в [Дорожной карте на 2021-2023 гг.](https://msal.ru/upload/medialibrary/99a/Dorozhnaya-karta-SAE-_-TSentr-prevoskhodstva.pdf)\n\n"
                                          "Руководитель и контактное лицо Центра правовых исследований банкротства – кандидат юридических наук, доцент кафедры предпринимательского и корпоративного права, адвокат Галкин Сергей Сергеевич (SSGalkin@msal.ru).\n\n"
                                          "Информация о проводимых мероприятиях и проектах САЕ размещается [в кафедральном Telegram-канале.](https://t.me/businesslawmsal)\n\n"
                                          "Для вступления в САЕ необходимо [заполнить анкету](https://disk.yandex.ru/i/vI0HhZPR6JCLWw) и направить ее на почту руководителю САЕ.", reply_markup=returnMarkup, parse_mode="Markdown",disable_web_page_preview=True)
    elif (message.text == "Школа предпринимателя"):
        returnMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        returnToMenuBtn = types.KeyboardButton("Вернуться в меню 🗂️")
        returnMarkup.row(returnToMenuBtn)
        bot.send_message(message.chat.id, "САЕ «Школа предпринимателя» создана 23 декабря 2021 г. на базе кафедры предпринимательского и корпоративного права, а также Института бизнес-права Университета имени О.Е. Кутафина (МГЮА).\n\n"
                                          "Деятельность САЕ «Школа предпринимателя» заключается в содействии обучающимся в создании и реализации стартап-проектов, взаимодействии с экспертами в области предпринимательства: выпускниками успешно запустившими стартапы, представителями юридического сообщества, включая государственных служащих и представителей институтов развития, развитии предпринимательских компетенций, а также проектной деятельности в сфере молодежного предпринимательства.\n\n"
                                          "Реализованные проекты:\n"
                                          "– [Программа повышения квалификации «Бизнес и право»](https://idop.msal.ru/businessandlaw) (29 июня–6 июля 2022 г., 20 ак. часов, 250+ слушателей);\n"
                                          "– [I Молодежный юридический слет «Кутафинский бизнес-форум»](https://msal.ru/news/-kutafinskiy-biznes-forum-v-mgyua-vspominaem-kak-eto-bylo/) (23-25 сентября 2022 г., загородный комплекс Университета имени О.Е. Кутафина (МГЮА) «Малюшина дача», 150+ участников);\n"
                                          "– [II Молодежный юридический слет «Кутафинский бизнес-форум»](https://www.msal.ru/news/ii-kutafinskiy-biznes-forum-sostoyalsya-v-mgyua/) (22-24 сентября 2023 г., загородный комплекс Университета имени О.Е. Кутафина (МГЮА) «Малюшина дача», 150+ участников).\n\n"
                                          "– [III Молодежный юридический слет «Кутафинский бизнес-форум»](https://msal.ru/news/v-zagorodnom-komplekse-malyushina-dacha-universiteta-imeni-o-e-kutafina-mgyua-sostoyalsya-iii-kutafi/) (22-22 сентября 2024 г., загородный комплекс Университета имени О.Е. Кутафина (МГЮА) «Малюшина дача», 130+ участников).\n\n"
                                          "– [Онлайн-курс «Правовое сопровождение технологичесого предпринимательства».](https://market.msal.ru/catalog/onlayn-kursy/pravovoe-soprovozhdenie-tekhnologicheskogo-predprinimatelstva/)\n\n"
                                          "В работе САЕ «Школа предпринимателя» принимают участие:\n"
                                          "– Петраков Андрей Юрьевич – кандидат юридических наук, директор Института бизнес-права, заместитель заведующего кафедрой предпринимательского и корпоративного права;\n"
                                          "– Оборов Александр Сергеевич – кандидат юридических наук, заместитель директора Института бизнес-права, доцент кафедры предпринимательского и корпоративного права;\n"
                                          "– Фролова Екатерина Константиновна – кандидат юридических наук, доцент кафедры предпринимательского и корпоративного права\n"
                                          "– Гузий Дмитрий Александрович – кандидат юридических наук, преподаватель кафедры гражданского и административного судопроизводства;\n"
                                          "– Вепринский Даниил Владимирович – аспирант кафедры предпринимательского и корпоративного права, председатель Совета молодых ученых Университета имени О.Е. Кутафина (МГЮА);\n"
                                          "Общая информация о планируемой деятельности САЕ «Школа предпринимателя» изложена [в Дорожной карте на 2020-2023 гг.](https://disk.yandex.ru/i/SjFugKNHi0-m4g)\n\n"
                                          "Руководитель САЕ «Школа предпринимателя» – кандидат юридических наук, преподаватель кафедры предпринимательского и корпоративного права Сайдашева Мария Владимировна (MVSAJDASHEVA@msal.ru).\n\n"
                                          "Информация о проводимых мероприятиях и проектах САЕ размещается [в кафедральном Telegram-канале](https://t.me/businesslawmsal), [собственном Telegram-канале](https://t.me/business_msal) и [группе VK.](https://vk.com/msalschoolofbusiness)\n\n"
                                          "Для вступления в САЕ необходимо [заполнить анкету](https://disk.yandex.ru/i/vI0HhZPR6JCLWw) и направить ее на почту [TABDULKADIROV@msal.ru](TABDULKADIROV@msal.ru)", reply_markup=returnMarkup, parse_mode="Markdown",disable_web_page_preview=True)
    elif (message.text == "Назад"):
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)

def choosePerson(message):
    person = message.text
    if (person == "Директор"):
        with open("Director.jpg", "rb") as photo:
            label = "Петраков Андрей Юрьевич\n\nДиректор Института бизнес-права, кандидат юридических наук, заместитель заведующего кафедрой предпринимательского и корпоративного права, член экспертных советов при Комитетах Государственной Думы по молодёжной политике и малому и среднему предпринимательству, член Комиссии Московского регионального отделения Ассоциации юристов России по правовой защите малого и среднего бизнеса, председатель Московского отделения Российского союза молодых ученых.\n\n🚪Кабинет: 412а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 796\n\n📬Электронная почта:\npetrakov@msal.ru\n\n⏰График работы:\nПН-ЧТ 09:30–18:15\nПТ 9:30–17:00\n(обед 13:30–14:00)"
            mesg = bot.send_photo(message.chat.id, photo, caption = label)
        bot.register_next_step_handler(mesg, choosePerson)

    elif (person == "Заместитель директора"):
        with open("Assistant.jpg", "rb") as photo:
            label = "Оборов Александр Сергеевич\n\nЗаместитель директора Института бизнес-права, кандидат юридических наук, преподаватель кафедры предпринимательского и корпоративного права, член Комиссии Московского регионального отделения Ассоциации юристов России по правовой защите малого и среднего бизнеса.\n\n🚪Кабинет: 412а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 811\n\n📬Электронная почта:\nASOBOROV@msal.ru\n\n⏰График работы:\nПН-ЧТ 8:45–17:30\nПТ 8:45–16:15\n(обед 13:30–14:00)"
            mesg = bot.send_photo(message.chat.id, photo, caption = label)
        bot.register_next_step_handler(mesg, choosePerson)

    elif (person == "Секретарь"):
        with open("Secretary.jpg", "rb") as photo:
            label = "Романова Анастасия Андреевна\n\n🚪Кабинет: 412а \n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 798\n\n📬Электронная почта:\nAAROMANOVA@msal.ru\n\n⏰График работы:\nПН-ЧТ 9:30–18:15\nПТ 9:30–17:00\n(обед 13:30–14:00)"
            mesg = bot.send_photo(message.chat.id, photo, caption = label)
        bot.register_next_step_handler(mesg, choosePerson)
    elif (person == "Научный руководитель"):
        with open("Scientific_supervisor.jpeg", "rb") as photo:
            label = "Ершова Инна Владимировна\n\nНаучный руководитель Института бизнес-права, доктор юридических наук, профессор, первый проректор, заведующий кафедрой предпринимательского и корпоративного права, руководитель Центра компетенций «Бизнес-право», заслуженный юрист Российской Федерации, Почетный работник высшего профессионального образования Российской Федерации, Почетный работник юстиции России, член Научно-консультативного совета Верховного Суда Российской Федерации, член Совета Московского отделения Ассоциации юристов России."
            mesg = bot.send_photo(message.chat.id, photo, caption=label)
        bot.register_next_step_handler(mesg, choosePerson)
    elif (person == "Инспекторы"):
        with open("Inspector1.jpg", "rb") as photo:
            label = "Соловей Наталья Сергеевна\n\nИнспектор\n2 курс очной формы обучения бакалавриата.\n\n🚪Кабинет: 402а \n(г. Москва, наб. Шитова, д. 72)\n\n📞 Телефон:\n+7 (499) 244-88-88 доб. 812\n\n📬Электронная почта:\nNSSOLOVEJ@msal.ru\n\n⏰График работы:\nПН-ЧТ 8:45–17:30\nПТ 8:45–16:15\n(обед 13:30–14:00)"
            bot.send_photo(message.chat.id, photo, caption=label)
        with open("Inspector2.jpg", "rb") as photo:
            label = "Заева Юлия Викторовна\n\nИнспектор\n3 и 4 курсы очной формы обучения бакалавриата.\n\n🚪Кабинет: 410а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 2013 \n📬Электронная почта: YUVZAEVA@msal.ru\n\n⏰График работы:\nПН-ЧТ 8:45–17:30\nПТ 8:45–16:15\n(обед 13:30–14:00)"
            bot.send_photo(message.chat.id, photo, caption=label)
        with open("Inspector3.jpg", "rb") as photo:
            label = "Шафеева Екатерина Александровна\n\nИнспектор\n1 курс очной формы обучения бакалавриата.\n\n🚪Кабинет: 402а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 809\n📬Электронная почта: EASHAFEEVA@msal.ru\n\n⏰График работы:\nПН, ВТ, ЧТ 8:45–17:30\nПТ 8:45–16:15\n(обед 13:30–14:00)\nСР 12:00–20:45\n(обед 16:30–17:00)"
            bot.send_photo(message.chat.id, photo, caption=label)
        with open("Inspector5.jpg", "rb") as photo:
            label = "Архарова Рамиля Наилевна\n\nИнспектор\nвсе курсы очно-заочной формы обучения бакалавриата и заочная форма обучения магистратуры (магистерские программы «Правовое сопровождение бизнеса (бизнес-юрист)» и «Корпоративное право»).\n\n🚪Кабинет: 502а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 0245\n📬Электронная почта: RNARHAROVA@msal.ru\n\n⏰График работы:\nПН, СР 8:45–17:30\nПТ 8:45–16:15\n(обед 13:30–14:00)\nВТ, ЧТ 12:00–20:45\n(обед 16:30–17:00)"
            bot.send_photo(message.chat.id, photo, caption=label)
        with open("Inspector4.jpg", "rb") as photo:
            label = "Алексашкина Людмила  Васильевна\n\nИнспектор\nочная и очно-заочная формы обучения магистратуры (магистерские программы «Правовое сопровождение\nбизнеса (бизнес-юрист)» и «Корпоративное право»).\n\n🚪Кабинет: 411а\n(г. Москва, наб. Шитова, д. 72)\n\n📞Телефон:\n+7 (499) 244-88-88 доб. 2097\n📬Электронная почта: LVALEKSASHKINA@msal.ru\n\n⏰График работы:\nПН-ЧТ 12:00–20:45\nПТ 12:00–19:30\n(обед 14:30–15:00)"
            mesg = bot.send_photo(message.chat.id, photo, caption=label)
        bot.register_next_step_handler(mesg, choosePerson)
    elif (person == "Главное меню 🗂️"):
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)


def chooseType(message):
    global typeOfRef, popravka4
    typeOfRef = message.text
    if typeOfRef == "↗️":
        pass
    elif typeOfRef == "↖️":
        pass
    elif (message.text == "Оформить справку 📋"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference2Button = types.InlineKeyboardButton("Офoрмить справку", url = "https://mfc.msal.ru/")# (o) (Оф_рмить справку)
        newsMarkup.row(reference2Button)
        mesg = bot.send_message(message.chat.id, "👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, chooseType) #fixed chooseType instead chooseAction
        return 
    elif typeOfRef == "Назад":
        menuKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        constitutionButton = types.KeyboardButton("Конституция России 🇷🇺")
        freqaskqueButton = types.KeyboardButton("Часто задаваемые вопросы 📄")
        orderRefButton = types.KeyboardButton("МФЦ")
        administrationButton = types.KeyboardButton("Руководство института")
        brochureButton = types.KeyboardButton("Брошюра об институте")
        studentsButton = types.KeyboardButton("Студентам 🧑‍🎓")
        applicantButton = types.KeyboardButton("Абитуриентам 🤵")
        mastersdegreeprogramsButton = types.KeyboardButton("Образовательные программы")
        newsButton = types.KeyboardButton("Новости 📰")
        contactsButton = types.KeyboardButton("Контакты 👥")
        startupButton = types.KeyboardButton("Программа «Стартап как ВКР»")
        CAEButton = types.KeyboardButton("Стратегические академические единицы")
        business_forumButton = types.KeyboardButton("Кутафинский бизнес-форум")
        menuKeyboard.row(administrationButton)
        menuKeyboard.row(brochureButton)
        menuKeyboard.row(newsButton, contactsButton)
        menuKeyboard.row(freqaskqueButton)
        menuKeyboard.row(orderRefButton)
        menuKeyboard.row(studentsButton, applicantButton)
        menuKeyboard.row(mastersdegreeprogramsButton)
        menuKeyboard.row(startupButton)
        menuKeyboard.row(CAEButton)
        menuKeyboard.row(business_forumButton)
        menuKeyboard.row(constitutionButton)
        mesg = bot.send_message(message.chat.id, "Выберите нужное действие:", reply_markup=menuKeyboard)
        bot.register_next_step_handler(mesg, chooseAction)
        return
    else:
        typeOfRefKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        referenceButton = types.KeyboardButton("Оформить справку 📋")
        defaultRefButton = types.KeyboardButton("↗️")
        refCallButton = types.KeyboardButton("↖️")
        typeOfRefKeyboard.row(referenceButton)
        typeOfRefKeyboard.row(defaultRefButton, refCallButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺", reply_markup=typeOfRefKeyboard)
        bot.register_next_step_handler(mesg, chooseType)
        return
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    backButton = types.KeyboardButton("Назад")
    markup.row(backButton)
    mesg = bot.send_message(message.chat.id,"Пожалуйста, отправьте своё имя в дательном падеже\n\nНапример:\n👉 Кому? Иванову Ивану Ивановичу\n👉 Кому? Ивановой Анастасии Игоревне",reply_markup=markup)
    bot.register_next_step_handler(mesg, takeFullName)
def takeFullName(message):
    global FullName
    FullName = message.text
    if message.text == "Назад":
        typeOfRefKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        referenceButton = types.KeyboardButton("Оформить справку 📋")
        defaultRefButton = types.KeyboardButton("↗️")
        refCallButton = types.KeyboardButton("↖️")
        backButton = types.KeyboardButton("Назад")
        typeOfRefKeyboard.row(referenceButton)
        typeOfRefKeyboard.row(defaultRefButton, refCallButton)
        typeOfRefKeyboard.row(backButton)
        mesg = bot.send_message(message.chat.id,"Выберите форму справки:\n\n•Справка об обучении: не требует дополнительных документов.\n\n•Справка-вызов: оформляется только при наличии справки с места работы и заявления; оригинал справки-вызова можно забрать только при предоставлении оригинала справки с места работы и оригинала заявления.",reply_markup=typeOfRefKeyboard)
        bot.register_next_step_handler(mesg, chooseType)
        return
    chooseGenderKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    maleButton = types.KeyboardButton("Мужской")
    femaleButton = types.KeyboardButton("Женский")
    backButton = types.KeyboardButton("Назад")
    chooseGenderKeyboard.add(maleButton)
    chooseGenderKeyboard.add(femaleButton)
    chooseGenderKeyboard.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите пол:", reply_markup=chooseGenderKeyboard)
    bot.register_next_step_handler(mesg, selectGender)
def selectGender(message):
    global gender, popravka1, popravka2, popravka3, popravka7
    gender = message.text
    if gender == "Мужской":
        popravka1 = "он"
        popravka2 = "обучающимся"
        popravka3 = "Зачислен"
        popravka7 = "обучающемуся"
    elif gender == "Женский":
        popravka1 = "она"
        popravka2 = "обучающейся"
        popravka3 = "Зачислена"
        popravka7 = "обучающейся"
    elif message.text == "Назад":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        backButton = types.KeyboardButton("Назад")
        markup.row(backButton)
        mesg = bot.send_message(message.chat.id,"Пожалуйста, отправьте своё имя в дательном падеже\n\nНапример:\n👉 Кому? Иванову Ивану Ивановичу\n👉 Кому? Ивановой Анастасии Игоревне",reply_markup=markup)
        bot.register_next_step_handler(mesg, takeFullName)
        return
    else:
        chooseGenderKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        maleButton = types.KeyboardButton("Мужской")
        femaleButton = types.KeyboardButton("Женский")
        chooseGenderKeyboard.add(maleButton)
        chooseGenderKeyboard.add(femaleButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺", reply_markup=chooseGenderKeyboard)
        bot.register_next_step_handler(mesg, selectGender)
        return
    chooseLevelKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    bachelorButton = types.KeyboardButton("Бакалавриат")
    magistracyButton = types.KeyboardButton("Магистратура")
    backButton = types.KeyboardButton("Назад")
    chooseLevelKeyboard.add(bachelorButton)
    chooseLevelKeyboard.add(magistracyButton)
    chooseLevelKeyboard.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите уровень получаемого образования:",reply_markup=chooseLevelKeyboard)
    bot.register_next_step_handler(mesg, selectLevelOfEducation)
def selectLevelOfEducation(message):
    global levelOfEducation
    levelOfEducation = message.text
    if levelOfEducation == "Бакалавриат":
        levelOfEducation = "бакалавриата"
    elif levelOfEducation == "Магистратура":
        levelOfEducation = "магистратуры"
    elif message.text == "Назад":
        chooseGenderKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        maleButton = types.KeyboardButton("Мужской")
        femaleButton = types.KeyboardButton("Женский")
        backButton = types.KeyboardButton("Назад")
        chooseGenderKeyboard.add(maleButton)
        chooseGenderKeyboard.add(femaleButton)
        chooseGenderKeyboard.row(backButton)
        mesg = bot.send_message(message.chat.id, "Выберите пол:", reply_markup=chooseGenderKeyboard)
        bot.register_next_step_handler(mesg, selectGender)
        return
    else:
        chooseLevelKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        bachelorButton = types.KeyboardButton("Бакалавриат")
        magistracyButton = types.KeyboardButton("Магистратура")
        chooseLevelKeyboard.add(bachelorButton)
        chooseLevelKeyboard.add(magistracyButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺",reply_markup=chooseLevelKeyboard)
        bot.register_next_step_handler(mesg, selectLevelOfEducation)
        return
    chooseFormOfEducationKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    FullTimeStudyButton = types.KeyboardButton("Очная")
    PartTimeStudyButton = types.KeyboardButton("Заочная")
    CombinedStudyButton = types.KeyboardButton("Очно-заочная")
    backButton = types.KeyboardButton("Назад")
    if levelOfEducation == "бакалавриата":
        chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
        chooseFormOfEducationKeyboard.add(CombinedStudyButton)
    elif levelOfEducation == "магистратуры":
        chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
        chooseFormOfEducationKeyboard.add(PartTimeStudyButton)
        chooseFormOfEducationKeyboard.add(CombinedStudyButton)
    chooseFormOfEducationKeyboard.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите форму обучения:", reply_markup=chooseFormOfEducationKeyboard)
    bot.register_next_step_handler(mesg, selectFormOfEducation)



    # chooseTypeOfStudy = types.ReplyKeyboardMarkup(resize_keyboard=True)
    # freeTypeButton = types.KeyboardButton("Бюджет")
    # payingTypeButton = types.KeyboardButton("Коммерция")
    # chooseTypeOfStudy.add(freeTypeButton)
    # chooseTypeOfStudy.add(payingTypeButton)
    # mesg = bot.send_message(message.chat.id, "Выберите основу обучения:", reply_markup=chooseTypeOfStudy)
    # bot.register_next_step_handler(mesg, takeYearOfAdmission)

    # chooseCourseKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    # firstCourseButton = types.KeyboardButton("1")
    # secondCourseButton = types.KeyboardButton("2")
    # thirdCourseButton = types.KeyboardButton("3")
    # fourthCourseButton = types.KeyboardButton("4")
    # chooseCourseKeyboard.row(firstCourseButton,secondCourseButton,thirdCourseButton,fourthCourseButton)
    # mesg = bot.send_message(message.chat.id, "Выберите курс:", reply_markup=chooseCourseKeyboard)
    # bot.register_next_step_handler(mesg, selectFormOfEducation)

def selectFormOfEducation(message):
    global formOfEducation
    formOfEducation = message.text
    if formOfEducation == "Очная":
        formOfEducation = "очной"
    elif formOfEducation == "Заочная":
        formOfEducation = "заочной"
    elif formOfEducation == "Очно-заочная":
        formOfEducation = "очно-заочной"
    elif message.text == "Назад":
        chooseLevelKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        bachelorButton = types.KeyboardButton("Бакалавриат")
        magistracyButton = types.KeyboardButton("Магистратура")
        backButton = types.KeyboardButton("Назад")
        chooseLevelKeyboard.add(bachelorButton)
        chooseLevelKeyboard.add(magistracyButton)
        chooseLevelKeyboard.row(backButton)
        mesg = bot.send_message(message.chat.id, "Выберите уровень получаемого образования:",reply_markup=chooseLevelKeyboard)
        bot.register_next_step_handler(mesg, selectLevelOfEducation)
        return
    else:
        chooseFormOfEducationKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        FullTimeStudyButton = types.KeyboardButton("Очная")
        PartTimeStudyButton = types.KeyboardButton("Заочная")
        CombinedStudyButton = types.KeyboardButton("Очно-заочная")
        if levelOfEducation == "бакалавриата":
            chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
            chooseFormOfEducationKeyboard.add(CombinedStudyButton)
        elif levelOfEducation == "магистратуры":
            chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
            chooseFormOfEducationKeyboard.add(PartTimeStudyButton)
            chooseFormOfEducationKeyboard.add(CombinedStudyButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺",reply_markup=chooseFormOfEducationKeyboard)
        bot.register_next_step_handler(mesg, selectFormOfEducation)
        return
    chooseCourseKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    firstCourseButton = types.KeyboardButton("1")
    secondCourseButton = types.KeyboardButton("2")
    thirdCourseButton = types.KeyboardButton("3")
    fourthCourseButton = types.KeyboardButton("4")
    fifthCourseButton = types.KeyboardButton("5")
    backButton = types.KeyboardButton("Назад")
    if levelOfEducation == "бакалавриата":
        if formOfEducation == "очной":
            chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton)
        elif formOfEducation == "очно-заочной":
            chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton, fifthCourseButton)
    elif levelOfEducation == "магистратуры":
        if formOfEducation == "очной":
            chooseCourseKeyboard.row(firstCourseButton, secondCourseButton)
        elif formOfEducation == "очно-заочной" or formOfEducation == "заочной":
            chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton)
    chooseCourseKeyboard.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите курс:", reply_markup=chooseCourseKeyboard)
    bot.register_next_step_handler(mesg, selectCourse)
def selectCourse(message):
    global course
    course = message.text
    if course == "1" or course == "2" or course == "3" or course == "4" or course == "5":
        pass
    elif message.text == "Назад":
        chooseFormOfEducationKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        FullTimeStudyButton = types.KeyboardButton("Очная")
        PartTimeStudyButton = types.KeyboardButton("Заочная")
        CombinedStudyButton = types.KeyboardButton("Очно-заочная")
        backButton = types.KeyboardButton("Назад")
        if levelOfEducation == "бакалавриата":
            chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
            chooseFormOfEducationKeyboard.add(CombinedStudyButton)
        elif levelOfEducation == "магистратуры":
            chooseFormOfEducationKeyboard.add(FullTimeStudyButton)
            chooseFormOfEducationKeyboard.add(PartTimeStudyButton)
            chooseFormOfEducationKeyboard.add(CombinedStudyButton)
        chooseFormOfEducationKeyboard.row(backButton)
        mesg = bot.send_message(message.chat.id, "Выберите форму обучения:", reply_markup=chooseFormOfEducationKeyboard)
        bot.register_next_step_handler(mesg, selectFormOfEducation)
        return
    else:
        chooseCourseKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        firstCourseButton = types.KeyboardButton("1")
        secondCourseButton = types.KeyboardButton("2")
        thirdCourseButton = types.KeyboardButton("3")
        fourthCourseButton = types.KeyboardButton("4")
        fifthCourseButton = types.KeyboardButton("5")
        if levelOfEducation == "бакалавриата":
            if formOfEducation == "очной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton)
            elif formOfEducation == "очно-заочной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton,fifthCourseButton)

        elif levelOfEducation == "магистратуры":
            if formOfEducation == "очной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton)
            elif formOfEducation == "очно-заочной" or formOfEducation == "заочной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺",reply_markup=chooseCourseKeyboard)
        bot.register_next_step_handler(mesg, selectCourse)
        return
    chooseTypeOfStudy = types.ReplyKeyboardMarkup(resize_keyboard=True)
    freeTypeButton = types.KeyboardButton("Бюджет")
    payingTypeButton = types.KeyboardButton("Коммерция")
    backButton = types.KeyboardButton("Назад")
    chooseTypeOfStudy.add(freeTypeButton)
    chooseTypeOfStudy.add(payingTypeButton)
    chooseTypeOfStudy.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите основу обучения:", reply_markup=chooseTypeOfStudy)
    bot.register_next_step_handler(mesg, takeYearOfAdmission)
def takeYearOfAdmission(message):
    global typeOfEducation
    typeOfEducation = message.text
    if typeOfEducation == "Бюджет":
        typeOfEducation = "бюджетной"
    elif typeOfEducation == "Коммерция":
        typeOfEducation = "платной"
    elif message.text == "Назад":
        chooseCourseKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        firstCourseButton = types.KeyboardButton("1")
        secondCourseButton = types.KeyboardButton("2")
        thirdCourseButton = types.KeyboardButton("3")
        fourthCourseButton = types.KeyboardButton("4")
        fifthCourseButton = types.KeyboardButton("5")
        backButton = types.KeyboardButton("Назад")
        if levelOfEducation == "бакалавриата":
            if formOfEducation == "очной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton)
            elif formOfEducation == "очно-заочной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton, fourthCourseButton,
                                         fifthCourseButton)
        elif levelOfEducation == "магистратуры":
            if formOfEducation == "очной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton)
            elif formOfEducation == "очно-заочной" or formOfEducation == "заочной":
                chooseCourseKeyboard.row(firstCourseButton, secondCourseButton, thirdCourseButton)
        chooseCourseKeyboard.row(backButton)
        mesg = bot.send_message(message.chat.id, "Выберите курс:", reply_markup=chooseCourseKeyboard)
        bot.register_next_step_handler(mesg, selectCourse)
        return
    else:
        chooseTypeOfStudy = types.ReplyKeyboardMarkup(resize_keyboard=True)
        freeTypeButton = types.KeyboardButton("Бюджет")
        payingTypeButton = types.KeyboardButton("Коммерция")
        chooseTypeOfStudy.add(freeTypeButton)
        chooseTypeOfStudy.add(payingTypeButton)
        mesg = bot.send_message(message.chat.id, "Пожалуйста, воспользуйся специальными кнопками 🥺",reply_markup=chooseTypeOfStudy)
        bot.register_next_step_handler(mesg, takeYearOfAdmission)
        return
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    backButton = types.KeyboardButton("Назад")
    markup.row(backButton)
    mesg = bot.send_message(message.chat.id, "В каком году вы поступили в Университет?\n\nГод должен быть только в виде числа❗️\n\nНапример:\n👉 2023\n👉 2022",reply_markup=markup)
    bot.register_next_step_handler(mesg,chooseOrganization)
def chooseOrganization(message):
    global yearOfAdmission
    if message.text == "Назад":
        chooseTypeOfStudy = types.ReplyKeyboardMarkup(resize_keyboard=True)
        freeTypeButton = types.KeyboardButton("Бюджет")
        payingTypeButton = types.KeyboardButton("Коммерция")
        backButton = types.KeyboardButton("Назад")
        chooseTypeOfStudy.add(freeTypeButton)
        chooseTypeOfStudy.add(payingTypeButton)
        chooseTypeOfStudy.row(backButton)
        mesg = bot.send_message(message.chat.id, "Выберите основу обучения:", reply_markup=chooseTypeOfStudy)
        bot.register_next_step_handler(mesg, takeYearOfAdmission)
        return
    try:
        yearOfAdmission = int(message.text)
    except:
        bot.send_message(message.chat.id, "Некорректный ввод данных!")
        mesg = bot.send_message(message.chat.id,"В каком году вы поступили в Университет?\n\n️️Год должен быть только в виде числа❗️\n\nНапример:\n👉 2023\n👉 2022")
        bot.register_next_step_handler(mesg, chooseOrganization)
        return
    if typeOfRef == "↗️":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        backButton = types.KeyboardButton("Назад")
        markup.row(backButton)
        mesg = bot.send_message(message.chat.id,'Введите полное наименование организации, куда представляется справка\n\nНапример:\n👉Федеральное государственное автономное образовательное учреждение высшего образования "Московский государственный юридический университет имени О.Е. Кутафина (МГЮА)"',reply_markup=backButton)
        bot.register_next_step_handler(mesg, result)
        return
    elif typeOfRef == "↖️":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        backButton = types.KeyboardButton("Назад")
        markup.row(backButton)
        mesg = bot.send_message(message.chat.id,'Введите полное наименование организации, куда представляется справка\n\nНапример:\n👉Федеральное государственное автономное образовательное учреждение высшего образования "Московский государственный юридический университет имени О.Е. Кутафина (МГЮА)"',reply_markup=backButton)
        bot.register_next_step_handler(mesg, chooseFormOfAtestation)
        return
    elif (message.text == "Оформить справку 📋"):
        newsMarkup = types.InlineKeyboardMarkup()
        reference2Button = types.InlineKeyboardButton("Оформить...", url = "https://mfc.msal.ru/")
        newsMarkup.row(reference2Button)
        mesg = bot.send_message(message.chat.id, "👇", reply_markup=newsMarkup)
        bot.register_next_step_handler(mesg, chooseAction)
        return

def chooseFormOfAtestation(message):
    if message.text == "Назад":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        backButton = types.KeyboardButton("Назад")
        markup.row(backButton)
        mesg = bot.send_message(message.chat.id,"В каком году вы поступили в Университет?\n\nГод должен быть только в виде числа❗️\n\nНапример:\n👉 2023\n👉 2022",reply_markup=markup)
        bot.register_next_step_handler(mesg, chooseOrganization)
        return
    global organizaton
    organizaton = message.text
    formOfAtestationMarkup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    giaBtn = types.KeyboardButton("Государственная итоговая аттестациия")
    promezhBtn = types.KeyboardButton("Промежуточная аттестация")
    backButton = types.KeyboardButton("Назад")
    formOfAtestationMarkup.row(giaBtn)
    formOfAtestationMarkup.row(promezhBtn)
    formOfAtestationMarkup.row(backButton)
    mesg = bot.send_message(message.chat.id, "Выберите форму аттестации", reply_markup=formOfAtestationMarkup)
    bot.register_next_step_handler(mesg, result)
def result(message):
    if message.text == "Назад":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        backButton = types.KeyboardButton("Назад")
        markup.row(backButton)
        mesg = bot.send_message(message.chat.id,"В каком году вы поступили в Университет?\n\nГод должен быть только в виде числа❗️\n\nНапример:\n👉 2023\n👉 2022",reply_markup=markup)
        bot.register_next_step_handler(mesg, chooseOrganization)
        return

    global organizaton, popravka5
    dateOfGraduation = "_____________________________."
    if formOfEducation == "очной":
        dateOfGraduation = f"июль {yearOfAdmission + 4} года."
        dateOfAdmission = f"1 сентября {yearOfAdmission} года."
    elif formOfEducation == "заочной" or formOfEducation == "очно-заочной":
        dateOfAdmission = f"1 ноября {yearOfAdmission} года." #FIX

    if levelOfEducation == "бакалавриата":
        popravka5 = "40.03.01"
        popravka6 = "по программе бакалавриата по направлению подготовки (специальности) 40.03.01 «Юриспруденция» (квалификация (степень) «бакалавр»)."
        if formOfEducation == "очной":
            dateOfGraduation = f"31 августа {yearOfAdmission + 4} года."
            dateOfAdmission = f"1 сентября {yearOfAdmission} года."
        elif formOfEducation == "очно-заочной":
            dateOfGraduation = f"1 марта {yearOfAdmission + 5} года."
            dateOfAdmission = f"1 сентября {yearOfAdmission} года."
    elif levelOfEducation == "магистратуры":
        popravka5 = "40.04.01"
        popravka6 = "по образовательной программе высшего образования - программе магистратуры по направлению подготовки (специальности) 40.04.01 «Юриспруденция» (квалификация (степень) «магистр»)."
        if formOfEducation == "очной":
            dateOfGraduation = f"31 августа {yearOfAdmission + 2} года."
            dateOfAdmission = f"1 сентября {yearOfAdmission} года."
        elif formOfEducation == "очно-заочной" or formOfEducation == "заочной":
            dateOfGraduation = f"31 января {yearOfAdmission + 2} года."
            dateOfAdmission = f"1 ноября {yearOfAdmission} года."

    if typeOfRef == "↗️":
        organizaton = message.text
        document = docx.Document("Document_copy_def.docx")
        document.paragraphs[19].text = f"Выдана {FullName} в том, что {popravka1} является {popravka2} {course} курса {formOfEducation} формы обучения Институт бизнес-права ФГАОУ ВО «Московский государственный юридический университет имени О.Е. Кутафина (МГЮА)» и обучается по основной профессиональной образовательной программе высшего образования — программе {levelOfEducation} по направлению подготовки {popravka5} «Юриспруденция» (уровень {levelOfEducation})."
        document.paragraphs[20].text = f"Обучается на {typeOfEducation} основе."
        document.paragraphs[21].text = f"Начало обучения с {dateOfAdmission}"
        document.paragraphs[22].text = f"{popravka3} приказом от _______________ года № ____."
        document.paragraphs[23].text = f"Предполагаемый срок окончания обучения {dateOfGraduation}"  # FIX убрать год
        document.paragraphs[24].text = f"Справка выдана для представления в {organizaton}."
        document.save('Document_done.docx')
    elif typeOfRef == "↖️":
        formOfAtestation = message.text
        if formOfAtestation == "Государственная итоговая аттестациия":
            formOfAtestation = "государственной итоговой аттестации"
        else:
            formOfAtestation = "промежуточной аттестации"
        document = docx.Document("Document_copy_.docx")
        document.paragraphs[22].text = f'Работодателю {organizaton} в соответствии со  статьей  173  Трудового кодекса Российской Федерации {FullName}, {popravka7} на {formOfEducation}  форме обучения на  {course}  курсе, предоставляются гарантии и компенсации для прохождения {formOfAtestation}  с ________________________ по ________________________ продолжительностью ___ календарных дней'
        document.paragraphs[23].text = f"Федеральное государственное автономное образовательное учреждение высшего образования «Московский государственный юридический университет имени О.Е. Кутафина (МГЮА)» имеет свидетельство о государственной аккредитации, выданное Федеральной службой по надзору в сфере образования и науки №3550, выданное 16.04.2021 по образовательной программе высшего образования - {popravka6}"

        document.save('Document_done.docx')
    bot.send_message(message.chat.id,"Справка готова✨\nПожалуйста, отправьте её секретарю на [корпоративную почту](https://mail.msal.ru/):\n\nAAROMANOVA@msal\.ru\n\nВ теме письма необходимо указать вид справки❗️", parse_mode='MarkdownV2')
    f = open("Document_done.docx", "rb")
    bot.send_document(message.chat.id, f)
    f.close()
    finalKeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    returnButton = types.KeyboardButton("Вернуться в меню 🗂️")
    finalKeyboard.add(returnButton)
    mesg = bot.send_message(message.chat.id,"Для продолжения пользования ботом, нажмите на кнопку ниже 👇",reply_markup=finalKeyboard)
    bot.register_next_step_handler(mesg, menu)
#FIX меню после отправки сообщений
#


bot.infinity_polling()
